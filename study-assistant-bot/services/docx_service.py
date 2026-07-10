"""
DOCX Processing Service
Extract text from Word documents
"""

from typing import Optional, List
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from utils.logger import logger


class DOCXService:
    """Service for processing DOCX documents"""

    @staticmethod
    def extract_text(file_path: str) -> Optional[str]:
        """
        Extract all text from DOCX file

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text or None if error
        """
        try:
            file_path = Path(file_path)

            if not file_path.exists():
                logger.error(f"DOCX file not found: {file_path}")
                return None

            if not file_path.suffix.lower() == ".docx":
                logger.error(f"File is not a DOCX: {file_path}")
                return None

            doc = Document(file_path)
            text = ""

            # Extract from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            return text

        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return None

    @staticmethod
    def extract_paragraphs(file_path: str) -> List[str]:
        """
        Extract paragraphs as separate items

        Args:
            file_path: Path to DOCX file

        Returns:
            List of paragraphs
        """
        try:
            file_path = Path(file_path)

            if not file_path.exists():
                return []

            doc = Document(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            return paragraphs

        except Exception as e:
            logger.error(f"Error extracting paragraphs from DOCX: {e}")
            return []

    @staticmethod
    def extract_tables(file_path: str) -> List[List[List[str]]]:
        """
        Extract tables from DOCX

        Args:
            file_path: Path to DOCX file

        Returns:
            List of tables (each table is list of rows, each row is list of cells)
        """
        try:
            file_path = Path(file_path)

            if not file_path.exists():
                return []

            doc = Document(file_path)
            tables = []

            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            return tables

        except Exception as e:
            logger.error(f"Error extracting tables from DOCX: {e}")
            return []

    @staticmethod
    def extract_headings(file_path: str) -> List[dict]:
        """
        Extract headings with their levels

        Args:
            file_path: Path to DOCX file

        Returns:
            List of headings with levels
        """
        try:
            file_path = Path(file_path)

            if not file_path.exists():
                return []

            doc = Document(file_path)
            headings = []

            for para in doc.paragraphs:
                style_name = para.style.name

                if "Heading" in style_name:
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        level = 0

                    headings.append({
                        "text": para.text,
                        "level": level,
                        "style": style_name,
                    })

            return headings

        except Exception as e:
            logger.error(f"Error extracting headings from DOCX: {e}")
            return []

    @staticmethod
    def get_document_structure(file_path: str) -> dict:
        """
        Get document structure with text and formatting

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with document structure
        """
        try:
            file_path = Path(file_path)

            if not file_path.exists():
                return {"error": "File not found"}

            doc = Document(file_path)

            structure = {
                "title": "",
                "paragraphs": [],
                "tables": [],
                "sections": [],
            }

            # Try to get title from first paragraph if it's formatted as title
            if doc.paragraphs:
                first_para = doc.paragraphs[0]
                if "Title" in first_para.style.name:
                    structure["title"] = first_para.text

            # Extract all content
            current_section = {"content": []}

            for element in doc.paragraphs:
                para_info = {
                    "text": element.text,
                    "style": element.style.name,
                    "level": element.paragraph_format.outline_level,
                }

                if "Heading" in element.style.name:
                    if current_section["content"]:
                        structure["sections"].append(current_section)
                    current_section = {
                        "heading": element.text,
                        "content": [],
                    }
                else:
                    current_section["content"].append(para_info)

            if current_section["content"] or "heading" in current_section:
                structure["sections"].append(current_section)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    "index": table_idx,
                    "rows": [],
                }
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data["rows"].append(row_data)
                structure["tables"].append(table_data)

            return structure

        except Exception as e:
            logger.error(f"Error getting document structure: {e}")
            return {"error": str(e)}

    @staticmethod
    def is_valid_docx(file_path: str) -> bool:
        """
        Check if file is a valid DOCX

        Args:
            file_path: Path to file

        Returns:
            True if valid DOCX, False otherwise
        """
        try:
            file_path = Path(file_path)

            if not file_path.exists():
                return False

            if not file_path.suffix.lower() == ".docx":
                return False

            # Try to open and verify
            Document(file_path)
            return True

        except Exception:
            return False

    @staticmethod
    def get_document_info(file_path: str) -> dict:
        """
        Get DOCX document information

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with document info
        """
        try:
            file_path = Path(file_path)

            if not file_path.exists():
                return {"error": "File not found"}

            doc = Document(file_path)

            # Get core properties
            props = doc.core_properties

            info = {
                "file_size": file_path.stat().st_size,
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "paragraphs_with_text": len([p for p in doc.paragraphs if p.text.strip()]),
            }

            return info

        except Exception as e:
            logger.error(f"Error getting DOCX info: {e}")
            return {"error": str(e)}

    @staticmethod
    def extract_formatted_text(file_path: str) -> dict:
        """
        Extract text with formatting information

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with formatted text
        """
        try:
            file_path = Path(file_path)

            if not file_path.exists():
                return {}

            doc = Document(file_path)
            content = {"text": "", "paragraphs": []}

            for para in doc.paragraphs:
                if not para.text.strip():
                    continue

                para_info = {
                    "text": para.text,
                    "style": para.style.name,
                    "bold": False,
                    "italic": False,
                    "underline": False,
                }

                # Check for formatting in runs
                for run in para.runs:
                    if run.bold:
                        para_info["bold"] = True
                    if run.italic:
                        para_info["italic"] = True
                    if run.underline:
                        para_info["underline"] = True

                content["paragraphs"].append(para_info)
                content["text"] += para.text + "\n"

            return content

        except Exception as e:
            logger.error(f"Error extracting formatted text: {e}")
            return {}

    @staticmethod
    def clean_extracted_text(text: str) -> str:
        """
        Clean extracted text

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text
        """
        # Remove extra whitespace
        text = " ".join(text.split())

        # Remove multiple consecutive punctuation
        import re
        text = re.sub(r"\.{2,}", ".", text)
        text = re.sub(r"\s{2,}", " ", text)

        return text
