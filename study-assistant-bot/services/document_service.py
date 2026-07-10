"""
Document processing service for Study Assistant Bot.
Handles PDF, DOCX, and TXT file processing.
"""

import logging
import os
from typing import Optional, List, Tuple
from pathlib import Path
import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument
from config import Config

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for processing and extracting content from documents."""
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
    
    @staticmethod
    def validate_file(file_path: str) -> Tuple[bool, Optional[str]]:
        """Validate file existence and type."""
        try:
            path = Path(file_path)
            
            if not path.exists():
                return False, "File not found"
            
            if path.suffix.lower() not in DocumentService.SUPPORTED_EXTENSIONS:
                return False, f"Unsupported file type. Supported: {', '.join(DocumentService.SUPPORTED_EXTENSIONS)}"
            
            file_size_mb = path.stat().st_size / (1024 * 1024)
            if file_size_mb > Config.MAX_UPLOAD_SIZE_MB:
                return False, f"File too large. Max size: {Config.MAX_UPLOAD_SIZE_MB}MB"
            
            return True, None
        except Exception as e:
            logger.error(f"Error validating file: {e}")
            return False, str(e)
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Tuple[Optional[str], int]:
        """Extract text from PDF file."""
        try:
            text = ""
            page_count = 0
            
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # Fallback to PyMuPDF if pdfplumber fails
            if not text:
                doc = fitz.open(file_path)
                page_count = len(doc)
                for page in doc:
                    text += page.get_text()
                doc.close()
            
            return text if text else None, page_count
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return None, 0
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            
            return text if text else None
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return None
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> Optional[str]:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text if text else None
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                return text if text else None
            except Exception as e:
                logger.error(f"Error reading TXT file: {e}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}")
            return None
    
    @staticmethod
    def extract_text(file_path: str) -> Tuple[Optional[str], int]:
        """Extract text from any supported file format."""
        try:
            path = Path(file_path)
            extension = path.suffix.lower()
            
            if extension == ".pdf":
                return DocumentService.extract_text_from_pdf(file_path)
            elif extension == ".docx":
                text = DocumentService.extract_text_from_docx(file_path)
                return text, 1
            elif extension == ".txt":
                text = DocumentService.extract_text_from_txt(file_path)
                return text, 1
            else:
                logger.error(f"Unsupported file format: {extension}")
                return None, 0
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return None, 0
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """Split text into overlapping chunks."""
        try:
            chunk_size = chunk_size or Config.CHUNK_SIZE
            overlap = overlap or Config.CHUNK_OVERLAP
            
            if not text or len(text) == 0:
                return []
            
            chunks = []
            start = 0
            
            while start < len(text):
                end = start + chunk_size
                chunk = text[start:end]
                
                # Try to split at sentence boundary
                if end < len(text):
                    last_period = chunk.rfind('.')
                    if last_period > chunk_size * 0.8:  # If period is reasonably close
                        end = start + last_period + 1
                    else:
                        # Try to find a space
                        last_space = chunk.rfind(' ')
                        if last_space > chunk_size * 0.8:
                            end = start + last_space
                
                chunk = text[start:end].strip()
                if chunk:
                    chunks.append(chunk)
                
                start = end - overlap
            
            return chunks
        except Exception as e:
            logger.error(f"Error chunking text: {e}")
            return []
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize text."""
        try:
            # Remove extra whitespace
            text = ' '.join(text.split())
            # Remove special characters but keep essential punctuation
            text = ''.join(c for c in text if c.isalnum() or c in ' .,!?;:\'"()\n')
            return text.strip()
        except Exception as e:
            logger.error(f"Error cleaning text: {e}")
            return text
